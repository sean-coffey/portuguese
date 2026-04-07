from copy import deepcopy
import os

from docx import Document
from docx.shared import Inches
from app.config import TEMPLATES_DIR

def replace_placeholder_in_doc(doc, placeholder, value):
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholder_in_paragraph(paragraph, placeholder, value)

    # Replace inside tables too (important!)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, placeholder, value)


def replace_placeholder_in_paragraph(paragraph, placeholder, value):
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return

    new_text = full_text.replace(placeholder, value or "")
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    for run in paragraph.runs:
        run.text = ""

    paragraph.runs[0].text = new_text


def _render_blank_lines(cell, count: int):
    for _ in range(count):
        cell.add_paragraph("________________________________________\n")


def build_docx(
    items,
    output_path,
    worksheet_profile,
    title: str | None = None,
    subtitle: str | None = None,
    instructions: str | None = None,
):
    template_path = os.path.join(TEMPLATES_DIR, worksheet_profile.template_filename)
    doc = Document(template_path)

    final_title = title or worksheet_profile.label_pt

    replace_placeholder_in_doc(doc, "{{TITLE}}", final_title)
    replace_placeholder_in_doc(doc, "{{SUBTITLE}}", subtitle or "")
    replace_placeholder_in_doc(doc, "{{INSTRUCTIONS}}", instructions or "")

    if not doc.tables:
        raise RuntimeError("Template must contain at least one table.")

    table = doc.tables[0]

    if not table.rows:
        raise RuntimeError("Template table must contain at least one row.")

    template_row = table.rows[0]

    for item in items:
        new_row_element = deepcopy(template_row._element)
        table._element.append(new_row_element)
        row = table.rows[-1]

        left_cell = row.cells[0]

        phrase_text = item.normalized if worksheet_profile.include_portuguese else ""
        gloss_text = item.gloss_en if worksheet_profile.include_english_gloss else ""
        note_text = "(review suggested)" if worksheet_profile.include_teacher_note and getattr(item, "teacher_review", False) else ""
        question_text = item.question_pt if worksheet_profile.include_question else ""

        for paragraph in left_cell.paragraphs:
            replace_placeholder_in_paragraph(paragraph, "{{PHRASE}}", phrase_text)
            replace_placeholder_in_paragraph(paragraph, "{{GLOSS}}", gloss_text)
            replace_placeholder_in_paragraph(paragraph, "{{NOTE}}", note_text)
            replace_placeholder_in_paragraph(paragraph, "{{QUESTION}}", question_text)

        if worksheet_profile.layout_mode == "image_with_lines":
            _render_blank_lines(left_cell, worksheet_profile.blank_lines_count)

        if worksheet_profile.layout_mode == "question_answer" and worksheet_profile.blank_lines_count > 0:
            _render_blank_lines(left_cell, worksheet_profile.blank_lines_count)

        right_cell = row.cells[1]
        img_paragraph = right_cell.paragraphs[0]
        run = img_paragraph.add_run()
        run.add_picture(item.image_path, width=Inches(3))

    table._element.remove(template_row._element)

    doc.save(output_path)
    return output_path

def build_instructions_for_profile(worksheet_profile) -> str:
    if worksheet_profile.name == "image_with_lines":
        return "Observa a imagem e responde à pergunta. Depois escreve a tua resposta nas linhas."

    if worksheet_profile.name == "vocabulary_question":
        return "Lê a pergunta com atenção e responde de forma clara."

    if worksheet_profile.name == "grammar_question":
        return "Lê cada frase com atenção e responde à pergunta."

    if worksheet_profile.name == "pt_en":
        return "Observa a imagem e lê o vocabulário apresentado."

    return ""

def build_subtitle_from_learner_profile(learner_profile) -> str:
    if learner_profile is None:
        return ""

    level = learner_profile.cefr_level if getattr(learner_profile, "cefr_level", None) else ""

    age_min = getattr(learner_profile, "age_min", None)
    age_max = getattr(learner_profile, "age_max", None)

    age_text = ""
    if age_min and age_max:
        age_text = f"{age_min}–{age_max} anos"
    elif age_min:
        age_text = f"{age_min}+ anos"
    elif age_max:
        age_text = f"Até {age_max} anos"

    if level and age_text:
        return f"{level} • {age_text}"
    if level:
        return level
    return age_text